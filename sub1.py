try:
    import Tkinter as tk
    import ttk
except ImportError:
    import tkinter as tk
    import tkinter.ttk as ttk
from tkinter import *
global str
from PIL import ImageTk, Image
import os
import time
import re
import random
from os.path import abspath, join, dirname
import names
from nltk.corpus import wordnet
from nltk.tokenize import word_tokenize
from random import randint
import nltk.data
import docx

root= tk.Tk()
root.geometry("1560x768")
root.title("Mathword Problem")
root.configure(background="#01b9f5")

t=[0,1]
str3=""

def generateQuestion():
        
    
    item = ['Beg', 'Pen', 'Book', 'Box', 'Chocolate', 'Biscuit', 'Car', 'Press', 'Mango', 'Banana', 'Doll', 'Flower', 'Bread', 'Watch', 'apple', 'apricot', 'avocado', 'banana', 'blackcurrant', 'blackberry', 'blueberry', 'cherry', 'coconut', 'fig', 'toy', 'kiwi', 'lemon', 'lime', 'lychee', 'mango', 'nectarine', 'orange', 'papaya', 'peach', 'pear', 'pineapple', 'plum', 'quince', 'raspberry', 'strawberry', 'watermelon']
    z = (random.choice(item))
    subtration = ['originally', 'in the first', 'in the beginning', 'earlier', 'to begin with', 'primitively', 'at first', 'initially', 'incipiently']
    az1 = (random.choice(subtration))
    __title__ = 'names'
    __version__ = '0.2'
    __author__ = 'Trey Hunner'
    __license__ = 'MIT'

    full_path = lambda filename: abspath(join(dirname(__file__), filename))

    FILES = {
        'first:male': full_path('dist.male.first'),
        'first:female': full_path('dist.female.first'),
        'last': full_path('dist.all.last'),
    }


    def multiwordReplace(text, wordDic):
        """
        take a text and replace words that match a key in a dictionary with
        the associated value, return the changed text
        """
        rc = re.compile('|'.join(map(re.escape, wordDic)))

        def translate(match):
            return wordDic[match.group(0)]

        return rc.sub(translate, text)


    p1 = random.randint(10, 99)
    p2 = random.randint(10, 99)
    f1 = max(p1,p2)
    f2 = min(p1,p2)
    q = str(f1)
    x = str(f2)


    def get_name(filename):
        selected = random.random() * 90
        with open(filename) as name_file:
            for line in name_file:
                name, _, cummulative, _ = line.split()
                if float(cummulative) > selected:
                    return name


    def get_first_name(gender=None):
        if gender not in ('male', 'female'):
            gender = random.choice(('male', 'female'))
        return get_name(FILES['first:%s' % gender]).capitalize()


    def get_last_name():
        return get_name(FILES['last']).capitalize()


    def get_full_name(gender=None):
        return u"%s %s" % (get_first_name(gender), get_last_name())


    tn = (names.get_first_name())
    p = (names.get_first_name())

    str1 = "John had some marbels. Jim took 3 from him. Now John has 8 marbels. How many marbels John had earlier?"
    # the dictionary has target_word : replacement_word pairs
    # print (str1)
    wordDic = {
        'John': tn,
        'marbels': z,
        'Jim': p,
        '8': q,
        '3': x,
        'earlier' : az1}
    # call the function and get the changed text
    str2 = multiwordReplace(str1, wordDic)
    str3 = (str2)
    #print (str2)
    output = ""

    # Load the pretrained neural net
    tokenizer = nltk.data.load('tokenizers/punkt/english.pickle')

    # Tokenize the text
    tokenized = tokenizer.tokenize(str3)

    # Get the list of words from the entire text
    words = word_tokenize(str3)

    # Identify the parts of speech
    tagged = nltk.pos_tag(words)

    

    for i in range(0, len(words)):
        replacements = []

        # Only replace nouns with nouns, vowels with vowels etc.
        for syn in wordnet.synsets(words[i]):

            # Do not attempt to replace proper nouns or determiners
            if tagged[i][1] == 'NNP' or tagged[i][1] == 'DT':
                break

            # The tokenizer returns strings like NNP, VBP etc
            # but the wordnet synonyms has tags like .n.
            # So we extract the first character from NNP ie n
            # then we check if the dictionary word has a .n. or not
            word_type = tagged[i][1][0].lower()
            if syn.name().find("." + word_type + "."):
                # extract the word only
                r = syn.name()[0:syn.name().find(".")]
                replacements.append(r)

        if len(replacements) > 0:
            # Choose a random replacement
            replacement = replacements[randint(0, len(replacements) - 1)]
            output = output + " " + replacement
        else:
            # If no replacement could be found, then just use the
            # original word
            output = output + " " + words[i]

    strr = str3
    tn = [int(s) for s in strr.split() if s.isdigit()]
    t[0] = tn[0]
    t[1] = tn[1]
    return strr

def OnExit():
    root.destroy()
    
def OnHome():
    root.destroy()
    os.system("python Home.py")    

def OnHints():
     root.destroy()
     os.system("python add1_c.py")
    

class Base_Form(object):
    """Base class of all forms"""

    def __init__(self, widget_class, master, action, hidden_input, kw):
        self.action = action

        if hidden_input is None:
            self.hidden_input = dict()
        else:
            if not isinstance(hidden_input, dict):
                raise ValueError("'hidden_input' should be a dict")
                
            self.hidden_input = hidden_input

        kw["class"] = "Form"
        widget_class.__init__(self, master, **kw)

class Base_SubmitButton(object):
    """Base class of submit buttons"""

    def submit(self):
        form_widget = self
        while True:
            form_widget = form_widget.master
            if form_widget is None:
                raise Exception("No equation found")
            else:
                if form_widget.winfo_class() == "Form":
                    break

        if form_widget.action is None: return
        
        form_action = form_widget.action

        form_data = {}
        form_data.update(form_widget.hidden_input)

        # Applying list for python 2/3 compatibility. dict_values is a view in Python 3.
        list_of_widgets = list(form_widget.children.values())
        
        while True:
            try:
                widget = list_of_widgets.pop()
            except IndexError:
                break

            list_of_widgets.extend(list(widget.children.values()))

            if not hasattr(widget,"fieldname"): continue
            
            field_name = widget.fieldname
            Tk_class = widget.winfo_class()
            
            if Tk_class == "Entry" or Tk_class == "TEntry":
                field_value = widget.get()
            elif Tk_class == "Text":               
                field_value = widget.get("1.0",'end-1c')
            elif Tk_class == "TCombobox":
                field_value = widget.get()
            elif Tk_class == "Listbox":
                field_value = [widget.get(idx) for idx in widget.curselection()]
            
            else:
                continue

            form_data[field_name] = field_value

        form_action(form_data)

class Form_Frame(tk.Frame, Base_Form):
    def __init__(self, master, action=None, hidden_input=None, **kw):
        Base_Form.__init__(self, tk.Frame, master, action, hidden_input, kw)
        
class Form_TFrame(tk.Frame, Base_Form):
    def __init__(self, master, action=None, hidden_input=None, **kw):
        Base_Form.__init__(self, ttk.Frame, master, action, hidden_input, kw)
        
class Form_LabelFrame(tk.LabelFrame, Base_Form):
    def __init__(self, master, action=None, hidden_input=None, **kw):
        Base_Form.__init__(self, tk.LabelFrame, master, action, hidden_input, kw)
        
class Form_TLabelFrame(ttk.LabelFrame, Base_Form):
    def __init__(self, master, action=None, hidden_input=None, **kw):
        Base_Form.__init__(self, ttk.LabelFrame, master, action, hidden_input, kw)

Form = Form_Frame

class Submit_Button(tk.Button, Base_SubmitButton):
    def __init__(self, parent, *args, **kw):
        kw["command"] = self.submit
        tk.Button.__init__(self, parent, *args, **kw)

class Submit_TButton(ttk.Button, Base_SubmitButton):
    def __init__(self, parent, *args, **kw):
        kw["command"] = self.submit
        ttk.Button.__init__(self, parent, *args, **kw)

 

if __name__== "__main__":
    try:
        from Tkinter import Frame, Entry, Radiobutton, Checkbutton, Text, Listbox, Tk, Label, StringVar
        import tkMessageBox as messagebox
        from ttk import Combobox
        from Tkconstants import *
    except ImportError:
        from tkinter import Frame, Entry, Radiobutton, Checkbutton, Text, Listbox, Tk, Label, messagebox, StringVar
        from tkinter.ttk import Combobox
        from tkinter.constants import *

    import pprint


pp = pprint.PrettyPrinter(indent=4)

class MainWindow(tk.Frame):
    counter = 0

    def __init__(self, *args, **kwargs):
        tk.Frame.__init__(self, *args, **kwargs)
        #self.str3 = "Ram has 7 marbles. Shyam gave him 3 more. How many marbles does Ram have now?"
        self.t = t
        #var.set("Equation Form Of The Mathword Problem Is:\n"+str(t[0])+ "+"+ str(t[1])+"= ?")
        self.k=self.t[0]+self.t[1]
        self.label = tk.Label( self,font=('arial', 15, 'bold'), bd=16,bg="#01b9f5", fg="#012b74", text="Equation Form Of The Mathword Problem Is:\n"+"? - " +str(self.t[0])+ " = "+ str(self.t[1]),relief=RAISED)
        self.label.pack()
        
        self.button = tk.Button(self, font=('arial', 15, 'bold'), bd=16,bg="#ff9900", fg="#012b74", text="Click here for answer", 
                                command=self.create_window)
        self.button.pack(side="top")

    def create_window(self):
        self.counter += 1
        t = tk.Toplevel(self)
        l = tk.Label(t, font=('arial', 25, 'bold'), bd=16,bg="#01b9f5", fg="#012b74", text="Correct Answer is->"+str(abs(self.k)))
        l.pack(side="top", fill="both", expand=True, padx=100, pady=100)

    def answer(self):
        self.k=self.t[1]+self.t[0]
        return str(abs(self.k))


img2 = ImageTk.PhotoImage(Image.open("math1.jpg"))
panel = Label(root, image = img2, bg="#01b9f5")
panel.pack(side = "bottom", fill = "both", expand = "yes")

#top label for topic 
Label(root, font=('aldhabi',30,'bold'), text="      QUESTION GENERATOR FOR MATHWORD PROBLEM", fg="white", bg="#01b9f5", bd=10, anchor='w').pack(anchor=W, padx=(0,0))

##uper part of the page
objt = MainWindow(root)
form = Form(root, action =lambda data: messagebox.showinfo("Submitted_Result_Is:",pp.pformat(data)+"\n\nCorrect answer is: "+ objt.answer()),bg="#99ccff")
form.pack(expand=True, fill="both", ipadx=10, ipady=10)

    # It's possible to provide hidden data
    # form.hidden_input["hidden_var1"] = "value1"
    # form.hidden_input["hidden_var2"] = "value2"

Label(form, height=1, font=('arial', 12, 'bold'), text="Question:", bg="#99ccff", fg='blue4', bd=20, anchor='w').grid(row=1,column=0, sticky=E, pady=(0,0))

#function to show and print question
entry = Entry(form, font=('arial',10,'bold'), width=120, bd=13, bg="#99ccff", fg="#012b74")
entry.fieldname = "Question"
entry.grid(row=1,column=1, sticky =E+W)
entry.delete(0, END)   
str3 = generateQuestion()
entry.insert(0, (str3))

Label(form, font=('arial', 12, 'bold'), text="Answer write here:", bg="#99ccff", fg="#012b74", bd=10, anchor='w').grid(row=5,column=0, sticky=E, pady=(8,2))
text1 = Text(form,font=('arial', 12, 'bold'), height=2, fg="#012b74", bg="#99ccff", bd=10)
text1.fieldname = "Submitted Answer"
text1.grid(row=5,column=1, sticky =E+W)

def nextQuestion():
    str3 = generateQuestion()
    entry.delete(0,1000)
    text1.delete(1.0,END)
    entry.insert(0, (str3))

def OnGenQuest():
    doc = docx.Document()
    doc.add_heading('Data set of MATHWORD PROBLEM',0)
    for i in range(0,3000):
        str3 = generateQuestion()
        entry.delete(0,END)
        entry.insert(0,(str3))
        doc.add_paragraph('{\nIndex ' + str(i) + ':')
        doc.add_paragraph('Question : "' + str3 + '"' )
        doc.add_paragraph('Equation : ' + '" X = '+ str(t[1]) + ' + ' + str(t[0]) + '"')
        doc.add_paragraph('Answer : "' +  str(t[1] + t[0]) + '" \n}')
    doc.save('dataSetSub1.docx')

Label(form, font=('algerian', 15, 'bold'), text=" For Equation Click On Yes", bg="#b0c4de", fg="#ff9900", bd=10, anchor='w').grid(row=26,column=0, sticky=E, pady=(15,0))

#creat button for connecting pages
bExit = Button(root, text='Exit', font=('algerian', 12, 'bold'), bd=10,
width=10, command=lambda:OnExit())
bExit.pack(side="left")
bYes = Button(root, text='Yes', font=('algerian', 12, 'bold'), bd=10,fg="#ff9900",
width=10, command=lambda:funct())
bYes.pack(side="left")
bHome = Button(root, text='Home', font=('algerian', 12, 'bold'), bd=10,
width=10, command=lambda:OnHome())
bHome.pack(side="left")
bNext = Button(root, text='Next', font=('algerian', 12, 'bold'), bd=10,
width=10, command=lambda:nextQuestion())
bNext.pack(side="left")
bHints = Button(root, text='Hints', font=('algerian', 12, 'bold'), bd=10,
width=10, command=lambda:OnHints())
bHints.pack(side="left")
bGenData = Button(root, text='Generate Data Set', font=('algerian', 12, 'bold'), bd=10,fg="#ff9900",
width=20, command=lambda:OnGenQuest())
bGenData.pack(side="left")


"""
def equation():
    print("Equation Form Of The Mathword Problem Is:\n", t[0], "+", t[1], "= ?")
    k = t[0] + t[1]
    print("do you want to see the answer of the problem Yes/No?")
    answer = input()

    if answer == 'Yes':
       print("Answer is:", abs(k))
       quit
    elif answer == 'No':
       print("Wow Great!! You Are Genius...")
       quit
"""
frame = Frame(root)
frame.pack()

def funct():
    root = tk.Tk()
    main = MainWindow(root)
    root.geometry("1560x768")
    root.title("Mathword Problem")
    main.pack(side="top", fill="both", expand=True)

##creat gap
Label(form, font=('arial',12,'bold'),bg="#99ccff",bd=8, fg="#99ccff", text=" ").grid(row=24,column=0, sticky=E, pady=(8,0))

Submit_Button(form,font=('arial',15,'bold'),bg="#b0c4de",bd=10, fg="#ff9900", text="Submit").grid(row=26,column=1,sticky =E)


root.mainloop()
